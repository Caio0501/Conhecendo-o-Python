from docx import Document
from docx.shared import Inches

def criandoDoc():     
    documento = Document();
    documento.add_heading('python word doc',0);
    documento.add_heading('Pacotes ofício',level = 1)
    documento.add_paragraph('Microsoft offce, libre office, Google offce,');
    documento.add_heading('Processador de texto',level = 1);
    documento.add_paragraph('microsoft word, LibreOffice Writer,  google docs')
    documento.add_heading('Planilha eletrônica',level = 1);
    documento.add_paragraph('microsoft excel, LibreOffice Calc, google sheets')
    documento.add_heading('Planilha eletrônica',level = 1);
    documento.add_paragraph('microsoft excel, LibreOffice Calc , google sheets')
    documento.add_heading('Apresentação de slides',level = 1);
    documento.add_paragraph('microsoft powerpoint, LibreOffice Impress, google slides')
    documento.add_heading('SGBD',level = 1);
    documento.add_paragraph('microsoft access, libreoffice base')
    documento.add_heading('Editor de diagramas',level = 1);
    documento.add_paragraph('microsoft visio, libreoffice draw, google drawing')
    documento.add_heading('Criador de publicações',level = 1);
    documento.add_paragraph('microsoft publisher ');
    documento.add_heading('cliente p/ correio eletrônico',level = 1);
    documento.add_paragraph('microsoft outlook, gmail')
    documento.add_heading('Programas Utilitário acessórios');
    documento.add_heading('Calculadora',level = 1);
    documento.add_paragraph('Windows calculadora, ubuntu calculadora');
    documento.add_heading('Navegador web',level = 1);
    documento.add_paragraph('google, mozilla firefox, opera');
    documento.add_heading('Visualizador de imagem',level = 1);
    documento.add_paragraph('Visualizados de de fotos do windows, Visualizador de imagens Nomacs, Galeria de fotos android');
    documento.add_heading('Explorador de arquivos e diretórios',level = 1);
    documento.add_paragraph('Windows Explorer, Nautilus (ubuntu), Astro File Manager (android)')
    documento.add_heading('Visualizador de processos do sistema operacional',level = 1);
    documento.add_paragraph('Gerenciadores de tarefas (windows), terminal comando "top" (linux)')
    documento.add_heading('Aplicativo para acesso remoto',level = 1);
    documento.add_paragraph('Splashtop , mRemoteNG (windows), AirDroid, AndroidonWeb (android), Remmina')
    documento.add_heading('Firewall',level = 1);
    documento.add_paragraph('Firewall Ubuntu, Windows Firewall');
    documento.add_heading('Referencia');
    documento.add_paragraph('https://python-docx.readthedocs.io/en/latest/');
    documento.add_paragraph('https://www.youtube.com/watch?v=oVhfTFV1oUk');

    documento.save('caio.docx');

doc = criandoDoc()       

